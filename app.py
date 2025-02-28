import streamlit as st  
import pandas as pd  
import numpy as np  
import matplotlib.pyplot as plt  
from scipy.special import factorial  
from sklearn.linear_model import LogisticRegression  
from sklearn.ensemble import RandomForestClassifier  
from xgboost import XGBClassifier  
from sklearn.model_selection import cross_val_score  
from sklearn.svm import SVC  
from io import BytesIO  
from docx import Document  

# Fonction pour prédire avec le modèle Poisson  
@st.cache_data  
def poisson_prediction(goals):  
    probabilities = np.exp(-goals) * (goals ** np.arange(6)) / factorial(np.arange(6))  
    return probabilities  

# Fonction pour créer un document Word avec les résultats  
def create_doc(results):  
    doc = Document()  
    doc.add_heading('🏆 Analyse de Matchs de Football et Prédictions de Paris Sportifs', level=1)  
    doc.add_heading('⚽ Données des Équipes', level=2)  
    doc.add_paragraph(f"🏠 Équipe Domicile: {results.get('Équipe Domicile', 'Inconnu')}")  
    doc.add_paragraph(f"🏟️ Équipe Extérieure: {results.get('Équipe Extérieure', 'Inconnu')}")  
    doc.add_paragraph(f"⚽ Buts Prédit Domicile: {float(results.get('Buts Prédit Domicile', 0)):.2f}")  
    doc.add_paragraph(f"⚽ Buts Prédit Extérieur: {float(results.get('Buts Prédit Extérieur', 0)):.2f}")  
    doc.add_heading('📊 Probabilités des Modèles', level=2)  
    doc.add_paragraph(f"🏠 Probabilité Domicile: {float(results.get('Probabilité Domicile', 0)):.2f}")  
    doc.add_paragraph(f"🤝 Probabilité Nul: {float(results.get('Probabilité Nul', 0)):.2f}")  
    doc.add_paragraph(f"🏟️ Probabilité Extérieure: {float(results.get('Probabilité Extérieure', 0)):.2f}")  
    buffer = BytesIO()  
    doc.save(buffer)  
    buffer.seek(0)  
    return buffer  

# Fonction pour entraîner et prédire avec les modèles  
@st.cache_resource  
def train_models():  
    np.random.seed(42)  
    data_size = 500  # Réduire la taille des données pour l'entraînement  
    data = pd.DataFrame({  
        'home_goals': np.random.randint(0, 3, size=data_size),  
        'away_goals': np.random.randint(0, 3, size=data_size),  
        'home_xG': np.random.uniform(0, 2, size=data_size),  
        'away_xG': np.random.uniform(0, 2, size=data_size),  
        'home_encais': np.random.uniform(0, 2, size=data_size),  
        'away_encais': np.random.uniform(0, 2, size=data_size),  
        'home_victories': np.random.randint(0, 20, size=data_size),  
        'away_victories': np.random.randint(0, 20, size=data_size),  
        'home_goals_scored': np.random.randint(0, 50, size=data_size),  
        'away_goals_scored': np.random.randint(0, 50, size=data_size),  
        'home_xGA': np.random.uniform(0, 2, size=data_size),  
        'away_xGA': np.random.uniform(0, 2, size=data_size),  
        'home_tirs_par_match': np.random.randint(0, 30, size=data_size),  
        'away_tirs_par_match': np.random.randint(0, 30, size=data_size),  
        'home_passes_cles_par_match': np.random.randint(0, 50, size=data_size),  
        'away_passes_cles_par_match': np.random.randint(0, 50, size=data_size),  
        'home_tirs_cadres': np.random.randint(0, 15, size=data_size),  
        'away_tirs_cadres': np.random.randint(0, 15, size=data_size),  
        'home_tirs_concedes': np.random.randint(0, 30, size=data_size),  
        'away_tirs_concedes': np.random.randint(0, 30, size=data_size),  
        'home_duels_defensifs': np.random.randint(0, 100, size=data_size),  
        'away_duels_defensifs': np.random.randint(0, 100, size=data_size),  
        'home_possession': np.random.uniform(0, 100, size=data_size),  
        'away_possession': np.random.uniform(0, 100, size=data_size),  
        'home_passes_reussies': np.random.uniform(0, 100, size=data_size),  
        'away_passes_reussies': np.random.uniform(0, 100, size=data_size),  
        'home_touches_surface': np.random.randint(0, 300, size=data_size),  
        'away_touches_surface': np.random.randint(0, 300, size=data_size),  
        'home_forme_recente': np.random.randint(0, 15, size=data_size),  
        'away_forme_recente': np.random.randint(0, 15, size=data_size),  
        'result': np.random.choice([0, 1, 2], size=data_size)  
    })  

    X = data.drop(columns='result')  
    y = data['result']  

    log_reg = LogisticRegression(max_iter=100, C=0.5, solver='lbfgs')  
    rf = RandomForestClassifier(n_estimators=50, max_depth=5, random_state=42)  
    xgb = XGBClassifier(use_label_encoder=False, eval_metric='logloss', n_estimators=50, max_depth=5, learning_rate=0.1)  
    svm = SVC(probability=True)  

    log_reg.fit(X, y)  
    rf.fit(X, y)  
    xgb.fit(X, y)  
    svm.fit(X, y)  

    return log_reg, rf, xgb, svm  

# Vérifier si les modèles sont déjà chargés dans l'état de session  
if 'models' not in st.session_state:  
    st.session_state.models = train_models()  

# Assurez-vous que les modèles sont bien chargés  
if st.session_state.models is not None:  
    log_reg_model, rf_model, xgb_model, svm_model = st.session_state.models  
else:  
    st.error("⚠️ Les modèles n'ont pas pu être chargés.")  

# Fonction pour évaluer les modèles avec validation croisée K-Fold  
@st.cache_data  
def evaluate_models(X, y):  
    models = {  
        "Régression Logistique": LogisticRegression(max_iter=100, C=0.5, solver='lbfgs'),  
        "Random Forest": RandomForestClassifier(n_estimators=50, max_depth=5, random_state=42),  
        "XGBoost": XGBClassifier(use_label_encoder=False, eval_metric='logloss', n_estimators=50, max_depth=5, learning_rate=0.1)  
    }  
    
    results = {}  
    for name, model in models.items():  
        scores = cross_val_score(model, X, y, cv=3, scoring='accuracy')  # K=3 pour une évaluation plus précise  
        results[name] = scores.mean()  
    
    return results  

# Interface utilisateur  
st.title("🏆 Analyse de Matchs de Football et Prédictions de Paris Sportifs")  

# Saisie des données des équipes  
st.header("📋 Saisie des données des équipes")  

# Création de deux colonnes pour les équipes  
col1, col2 = st.columns(2)  

# Équipe à domicile  
with col1:  
    st.subheader("Équipe à Domicile")  
    home_team = st.text_input("🏠 Nom de l'équipe à domicile", value="Équipe A")  
    home_goals = st.number_input("⚽ Moyenne de buts marqués par match (domicile)", min_value=0.0, max_value=5.0, value=2.5)  
    home_xG = st.number_input("📈 xG (Expected Goals) (domicile)", min_value=0.0, max_value=5.0, value=2.0)  
    home_encais = st.number_input("🚫 Moyenne de buts encaissés par match (domicile)", min_value=0.0, max_value=5.0, value=1.0)  
    home_victories = st.number_input("🏆 Nombre de victoires à domicile", min_value=0, value=5)  
    home_goals_scored = st.number_input("⚽ Nombre de buts marqués à domicile", min_value=0, value=15)  
    home_xGA = st.number_input("📉 xGA (Expected Goals Against) (domicile)", min_value=0.0, max_value=5.0, value=1.5)  
    home_tirs_par_match = st.number_input("🔫 Nombres de tirs par match (domicile)", min_value=0.0, max_value=30.0, value=15.0)  
    home_passes_cles_par_match = st.number_input("📊 Nombres de passes clés par match (domicile)", min_value=0.0, max_value=50.0, value=10.0)  
    home_tirs_cadres = st.number_input("🎯 Tirs cadrés par match (domicile)", min_value=0.0, max_value=15.0, value=5.0)  
    home_tirs_concedes = st.number_input("🚫 Nombres de tirs concédés par match (domicile)", min_value=0.0, max_value=30.0, value=8.0)  
    home_duels_defensifs = st.number_input("🤼 Duels défensifs gagnés (domicile)", min_value=0.0, max_value=100.0, value=60.0)  
    home_possession = st.number_input("📊 Possession moyenne (%) (domicile)", min_value=0.0, max_value=100.0, value=55.0)  
    home_passes_reussies = st.number_input("✅ Passes réussies (%) par match (domicile)", min_value=0.0, max_value=100.0, value=80.0)  
    home_touches_surface = st.number_input("⚽ Balles touchées dans la surface adverse par match (domicile)", min_value=0.0, max_value=300.0, value=20.0)  
    home_forme_recente = st.number_input("📈 Forme récente (points sur les 5 derniers matchs) (domicile)", min_value=0, max_value=15, value=10)  

# Équipe à l'extérieur  
with col2:  
    st.subheader("Équipe à Extérieur")  
    away_team = st.text_input("🏟️ Nom de l'équipe à l'extérieur", value="Équipe B")  
    away_goals = st.number_input("⚽ Moyenne de buts marqués par match (extérieur)", min_value=0.0, max_value=5.0, value=1.5)  
    away_xG = st.number_input("📈 xG (Expected Goals) (extérieur)", min_value=0.0, max_value=5.0, value=1.8)  
    away_encais = st.number_input("🚫 Moyenne de buts encaissés par match (extérieur)", min_value=0.0, max_value=5.0, value=2.0)  
    away_victories = st.number_input("🏆 Nombre de victoires à l'extérieur", min_value=0, value=3)  
    away_goals_scored = st.number_input("⚽ Nombre de buts marqués à l'extérieur", min_value=0, value=10)  
    away_xGA = st.number_input("📉 xGA (Expected Goals Against) (extérieur)", min_value=0.0, max_value=5.0, value=1.5)  
    away_tirs_par_match = st.number_input("🔫 Nombres de tirs par match (extérieur)", min_value=0.0, max_value=30.0, value=12.0)  
    away_passes_cles_par_match = st.number_input("📊 Nombres de passes clés par match (extérieur)", min_value=0.0, max_value=50.0, value=8.0)  
    away_tirs_cadres = st.number_input("🎯 Tirs cadrés par match (extérieur)", min_value=0.0, max_value=15.0, value=4.0)  
    away_tirs_concedes = st.number_input("🚫 Nombres de tirs concédés par match (extérieur)", min_value=0.0, max_value=30.0, value=10.0)  
    away_duels_defensifs = st.number_input("🤼 Duels défensifs gagnés (extérieur)", min_value=0.0, max_value=100.0, value=55.0)  
    away_possession = st.number_input("📊 Possession moyenne (%) (extérieur)", min_value=0.0, max_value=100.0, value=50.0)  
    away_passes_reussies = st.number_input("✅ Passes réussies (%) (extérieur)", min_value=0.0, max_value=100.0, value=75.0)  
    away_touches_surface = st.number_input("⚽ Balles touchées dans la surface adverse par match (extérieur)", min_value=0.0, max_value=300.0, value=15.0)  
    away_forme_recente = st.number_input("📈 Forme récente (points sur les 5 derniers matchs) (extérieur)", min_value=0, max_value=15, value=8)  

# Saisie des cotes des bookmakers (non utilisées par les modèles)  
st.header("💰 Cotes des Équipes")  
odds_home = st.number_input("🏠 Cote pour l'équipe à domicile", min_value=1.0, value=1.8)  
odds_away = st.number_input("🏟️ Cote pour l'équipe à l'extérieur", min_value=1.0, value=2.2)  

# Calcul des probabilités implicites  
def calculate_implied_prob(odds):  
    return 1 / odds  

# Fonction pour afficher les graphiques des performances des équipes  
def plot_team_performance(home_stats, away_stats):  
    labels = ['Buts', 'xG', 'Buts Encaissés', 'Tirs', 'Passes Clés', 'Tirs Cadrés', 'Possession']  
    home_values = [home_stats['home_goals_scored'], home_stats['home_xG'], home_stats['home_encais'],   
                   home_stats['home_tirs_par_match'], home_stats['home_passes_cles_par_match'],   
                   home_stats['home_tirs_cadres'], home_stats['home_possession']]  
    away_values = [away_stats['away_goals_scored'], away_stats['away_xG'], away_stats['away_encais'],   
                   away_stats['away_tirs_par_match'], away_stats['away_passes_cles_par_match'],   
                   away_stats['away_tirs_cadres'], away_stats['away_possession']]  

    x = np.arange(len(labels))  
    width = 0.35  

    fig, ax = plt.subplots()  
    bars1 = ax.bar(x - width/2, home_values, width, label=home_team)  
    bars2 = ax.bar(x + width/2, away_values, width, label=away_team)  

    ax.set_ylabel('Valeurs')  
    ax.set_title('Comparaison des Performances des Équipes')  
    ax.set_xticks(x)  
    ax.set_xticklabels(labels)  
    ax.legend()  

    st.pyplot(fig)  

# Prédictions  
if st.button("🔍 Prédire les résultats"):  
    with st.spinner('Calcul des résultats...'):  
        try:  
            # Validation des entrées utilisateur  
            if home_goals < 0 or away_goals < 0:  
                st.error("⚠️ Les moyennes de buts ne peuvent pas être négatives.")  
            else:  
                # Évaluation des modèles avec validation croisée K-Fold  
                X = pd.DataFrame({  
                    'home_goals': np.random.randint(0, 3, size=500),  
                    'away_goals': np.random.randint(0, 3, size=500),  
                    'home_xG': np.random.uniform(0, 2, size=500),  
                    'away_xG': np.random.uniform(0, 2, size=500),  
                    'home_encais': np.random.uniform(0, 2, size=500),  
                    'away_encais': np.random.uniform(0, 2, size=500),  
                    'home_victories': np.random.randint(0, 20, size=500),  
                    'away_victories': np.random.randint(0, 20, size=500),  
                    'home_goals_scored': np.random.randint(0, 50, size=500),  
                    'away_goals_scored': np.random.randint(0, 50, size=500),  
                    'home_xGA': np.random.uniform(0, 2, size=500),  
                    'away_xGA': np.random.uniform(0, 2, size=500),  
                    'home_tirs_par_match': np.random.randint(0, 30, size=500),  
                    'away_tirs_par_match': np.random.randint(0, 30, size=500),  
                    'home_passes_cles_par_match': np.random.randint(0, 50, size=500),  
                    'away_passes_cles_par_match': np.random.randint(0, 50, size=500),  
                    'away_tirs_cadres': np.random.randint(0, 15, size=500),  
                    'home_tirs_concedes': np.random.randint(0, 30, size=500),  
                    'away_tirs_concedes': np.random.randint(0, 30, size=500),  
                    'home_duels_defensifs': np.random.randint(0, 100, size=500),  
                    'away_duels_defensifs': np.random.randint(0, 100, size=500),  
                    'home_possession': np.random.uniform(0, 100, size=500),  
                    'away_possession': np.random.uniform(0, 100, size=500),  
                    'home_passes_reussies': np.random.uniform(0, 100, size=500),  
                    'away_passes_reussies': np.random.uniform(0, 100, size=500),  
                    'home_touches_surface': np.random.randint(0, 300, size=500),  
                    'away_touches_surface': np.random.randint(0, 300, size=500),  
                    'home_forme_recente': np.random.randint(0, 15, size=500),  
                    'away_forme_recente': np.random.randint(0, 15, size=500)  
                })  
                y = np.random.choice([0, 1, 2], size=500)  # 0: Domicile, 1: Nul, 2: Extérieur  
                
                # Évaluation des modèles avec validation croisée K=3  
                results = evaluate_models(X, y)  
                st.write("📊 Résultats de la validation croisée K-Fold :", results)  

                # Calcul des buts prédit  
                home_goals_pred = home_goals + home_xG - away_encais  
                away_goals_pred = away_goals + away_xG - home_encais  

                # Calcul des probabilités avec le modèle de Poisson  
                home_probabilities = poisson_prediction(home_goals_pred)  
                away_probabilities = poisson_prediction(away_goals_pred)  

                # Formatage des résultats pour l'affichage  
                home_results = {i: home_probabilities[i] * 100 for i in range(len(home_probabilities))}  
                away_results = {i: away_probabilities[i] * 100 for i in range(len(away_probabilities))}  

                # Création d'un DataFrame pour les résultats de Poisson  
                poisson_results = pd.DataFrame({  
                    "Nombre de Buts": range(len(home_probabilities)),  
                    f"Probabilités {home_team} (%)": [f"{home_results[i]:.2f}" for i in range(len(home_results))],  
                    f"Probabilités {away_team} (%)": [f"{away_results[i]:.2f}" for i in range(len(away_results))]  
                })  

                # Affichage des résultats du modèle de Poisson  
                st.markdown("### Résultats du Modèle de Poisson")  
                st.dataframe(poisson_results, use_container_width=True)  

                # Détails sur chaque prédiction des modèles  
                st.markdown("### Détails des Prédictions des Modèles")  
                model_details = {  
                    "Modèle": ["Régression Logistique", "Random Forest", "XGBoost", "SVM"],  
                    "Probabilité Domicile (%)": [  
                        log_reg_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                       home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                       home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                       home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                       away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                       home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                       away_possession, home_passes_reussies, away_passes_reussies,  
                                                       home_touches_surface, away_touches_surface, home_forme_recente,  
                                                       away_forme_recente]])[0][0] * 100,  
                        rf_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                  home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                  home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                  home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                  away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                  home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                  away_possession, home_passes_reussies, away_passes_reussies,  
                                                  home_touches_surface, away_touches_surface, home_forme_recente,  
                                                  away_forme_recente]])[0][0] * 100,  
                        xgb_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                   home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                   home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                   home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                   away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                   home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                   away_possession, home_passes_reussies, away_passes_reussies,  
                                                   home_touches_surface, away_touches_surface, home_forme_recente,  
                                                   away_forme_recente]])[0][0] * 100,  
                        svm_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                   home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                   home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                   home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                   away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                   home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                   away_possession, home_passes_reussies, away_passes_reussies,  
                                                   home_touches_surface, away_touches_surface, home_forme_recente,  
                                                   away_forme_recente]])[0][0] * 100  
                    ],  
                    "Probabilité Nul (%)": [  
                        log_reg_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                       home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                       home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                       home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                       away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                       home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                       away_possession, home_passes_reussies, away_passes_reussies,  
                                                       home_touches_surface, away_touches_surface, home_forme_recente,  
                                                       away_forme_recente]])[0][1] * 100,  
                        rf_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                  home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                  home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                  home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                  away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                  home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                  away_possession, home_passes_reussies, away_passes_reussies,  
                                                  home_touches_surface, away_touches_surface, home_forme_recente,  
                                                  away_forme_recente]])[0][1] * 100,  
                        xgb_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                   home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                   home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                   home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                   away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                   home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                   away_possession, home_passes_reussies, away_passes_reussies,  
                                                   home_touches_surface, away_touches_surface, home_forme_recente,  
                                                   away_forme_recente]])[0][1] * 100,  
                        svm_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                   home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                   home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                   home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                   away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                   home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                   away_possession, home_passes_reussies, away_passes_reussies,  
                                                   home_touches_surface, away_touches_surface, home_forme_recente,  
                                                   away_forme_recente]])[0][1] * 100  
                    ],  
                    "Probabilité Extérieure (%)": [  
                        log_reg_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                       home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                       home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                       home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                       away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                       home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                       away_possession, home_passes_reussies, away_passes_reussies,  
                                                       home_touches_surface, away_touches_surface, home_forme_recente,  
                                                       away_forme_recente]])[0][2] * 100,  
                        rf_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                  home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                  home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                  home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                  away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                  home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                  away_possession, home_passes_reussies, away_passes_reussies,  
                                                  home_touches_surface, away_touches_surface, home_forme_recente,  
                                                  away_forme_recente]])[0][2] * 100,  
                        xgb_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                   home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                   home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                   home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                   away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                   home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                   away_possession, home_passes_reussies, away_passes_reussies,  
                                                   home_touches_surface, away_touches_surface, home_forme_recente,  
                                                   away_forme_recente]])[0][2] * 100,  
                        svm_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                   home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                   home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                   home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                   away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                   home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                   away_possession, home_passes_reussies, away_passes_reussies,  
                                                   home_touches_surface, away_touches_surface, home_forme_recente,  
                                                   away_forme_recente]])[0][2] * 100  
                    ]  
                }  
                model_details_df = pd.DataFrame(model_details)  
                st.dataframe(model_details_df, use_container_width=True)  

                # Comparaison des probabilités implicites et prédites  
                st.subheader("📊 Comparaison des Probabilités Implicites et Prédites")  
                implied_home_prob = calculate_implied_prob(odds_home)  
                implied_away_prob = calculate_implied_prob(odds_away)  
                implied_draw_prob = 1 - (implied_home_prob + implied_away_prob)  

                comparison_data = {  
                    "Type": ["Implicite Domicile", "Implicite Nul", "Implicite Extérieure",   
                             "Prédite Domicile", "Prédite Nul", "Prédite Extérieure"],  
                    "Probabilité (%)": [  
                        implied_home_prob * 100,  
                        implied_draw_prob * 100,  
                        implied_away_prob * 100,  
                        log_reg_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                       home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                       home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                       home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                       away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                       home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                       away_possession, home_passes_reussies, away_passes_reussies,  
                                                       home_touches_surface, away_touches_surface, home_forme_recente,  
                                                       away_forme_recente]])[0][0] * 100,  
                        log_reg_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                       home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                       home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                       home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                       away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                       home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                       away_possession, home_passes_reussies, away_passes_reussies,  
                                                       home_touches_surface, away_touches_surface, home_forme_recente,  
                                                       away_forme_recente]])[0][1] * 100,  
                        log_reg_model.predict_proba([[home_goals_pred, away_goals_pred, home_xG, away_xG, home_encais, away_encais,  
                                                       home_victories, away_victories, home_goals_scored, away_goals_scored,  
                                                       home_xGA, away_xGA, home_tirs_par_match, away_tirs_par_match,  
                                                       home_passes_cles_par_match, away_passes_cles_par_match, home_tirs_cadres,  
                                                       away_tirs_cadres, home_tirs_concedes, away_tirs_concedes,  
                                                       home_duels_defensifs, away_duels_defensifs, home_possession,  
                                                       away_possession, home_passes_reussies, away_passes_reussies,  
                                                       home_touches_surface, away_touches_surface, home_forme_recente,  
                                                       away_forme_recente]])[0][2] * 100  
                    ]  
                }  
                comparison_df = pd.DataFrame(comparison_data)  
                st.dataframe(comparison_df, use_container_width=True)  

                # Affichage des graphiques des performances des
